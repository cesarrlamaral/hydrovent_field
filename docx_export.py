"""
Exportação do pacote de submissão (HTML -> DOCX, tabelas em DOCX próprio,
carta de apresentação e arquivo suplementar .zip), usado só pelos geradores
de paper do Administrador em report.py. FORA do pacote git (ver .gitignore),
mesmo motivo de report.py/relatorios_admin.py: conteúdo específico do
manuscrito do autor, não parte do software genérico de simulação.

Depende de `python-docx` (pip install python-docx) — dependência SÓ desta
ferramenta interna do Administrador, deliberadamente NÃO adicionada a
requirements.txt (que documenta as dependências do software público).

Regras reais de arquivo suplementar usadas aqui (verificadas 2026-08-09):
  - PLOS ONE (journals.plos.org/plosone/s/supporting-information): qualquer
    tipo de arquivo aceito; nomeação obrigatória "S1 Fig", "S1 Table", "S1
    Appendix" etc., legendas ao final do manuscrito.
  - SAGE (sagepub.com/journals/.../supplemental-material-guidelines-for-authors):
    lista ampla de formatos aceitos (docx, csv, py, json, zip entre outros),
    mas "zip files should be avoided wherever possible" — SAGE prefere
    arquivos individuais, cada um com título/descrição próprios. O .zip que
    este módulo gera é para conveniência do autor (tudo num lugar só,
    revisão local); o README dentro dele documenta essa ressalva real.
  - Carta de apresentação da PLOS ONE (journals.plos.org/plosone/s/submission-guidelines):
    deve conter: contribuição do estudo, relação com a literatura, tipo de
    artigo, interações prévias com a PLOS, sugestão de editores acadêmicos,
    revisores a excluir; limite de 1 página; NUNCA pedir isenção de taxa ali.
  - Carta de apresentação da SAGE/Astrobiology: sem checklist formal
    publicado, mas com uma exigência real e específica (Writing assistance
    and third party submissions, docs/ASTROBIOLOGY_SAGE_AUTHOR_RULES.txt):
    qualquer assistência de terceiros na escrita/edição — o que inclui uso
    de IA — deve ser declarada tanto em Acknowledgments quanto na carta.
"""

from __future__ import annotations

import os
import re
import zipfile
from datetime import datetime
from html.parser import HTMLParser
from typing import List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image

# --------------------------------------------------------------------------
# Padrão de figura em tamanho/resolução de publicação
# --------------------------------------------------------------------------

MIN_FIGURE_WIDTH_CM = 20.0
MIN_FIGURE_DPI = 300
MIN_FIGURE_WIDTH_PX = int(round(MIN_FIGURE_WIDTH_CM / 2.54 * MIN_FIGURE_DPI))  # ~2362px


def enforce_matplotlib_min_size(fig) -> None:
    """Amplia (mantendo a proporção) um matplotlib.figure.Figure, se
    necessário, para que salvo a 300dpi tenha pelo menos MIN_FIGURE_WIDTH_CM
    de largura. Fontes/elementos ficam do mesmo tamanho ABSOLUTO (pontos),
    só a tela física cresce — comportamento padrão e correto para
    publicação (ver report.FigureManager.add)."""
    w_in, h_in = fig.get_size_inches()
    min_w_in = MIN_FIGURE_WIDTH_CM / 2.54
    if w_in < min_w_in:
        scale = min_w_in / w_in
        fig.set_size_inches(w_in * scale, h_in * scale)


def enforce_raster_min_size(fpath: str) -> bool:
    """Reamostra (Lanczos) uma imagem já salva em `fpath`, SE seus pixels
    reais não bastarem para 20cm a 300dpi, e regrava com metadado de dpi
    real e honesto (dpi implícito = pixels / (largura_alvo_pol)). Devolve
    True se a imagem foi reamostrada (nitidez REAL não aumenta — só o
    número de pixels; ver aviso equivalente em FigureManager.any_upsampled).
    Se a imagem já atende ao mínimo, só grava dpi=300 real (sem reamostrar)
    e devolve False."""
    with Image.open(fpath) as img:
        img = img.convert("RGB")
        w_px, h_px = img.size
        if w_px >= MIN_FIGURE_WIDTH_PX:
            img.save(fpath, dpi=(300, 300))
            return False
        scale = MIN_FIGURE_WIDTH_PX / w_px
        new_size = (MIN_FIGURE_WIDTH_PX, int(round(h_px * scale)))
        img = img.resize(new_size, Image.LANCZOS)
        img.save(fpath, dpi=(300, 300))
        return True


# --------------------------------------------------------------------------
# HTML (vocabulário controlado de report.py) -> DOCX
# --------------------------------------------------------------------------

class _DocxHTMLConverter(HTMLParser):
    """Converte o HTML gerado internamente por report.py (vocabulário
    limitado e conhecido: h1-h3, p, ul/li, b/strong, i/em, sup, sub, code,
    table/caption/tr/th/td, figure/figcaption/img, br) para um
    docx.Document já aberto. NÃO é um conversor HTML genérico — depende de
    report.py nunca emitir tags fora desse vocabulário."""

    _INLINE_TAGS = ("b", "strong", "i", "em", "sup", "sub", "code")

    def __init__(self, document: Document, figures_dir: Optional[str] = None):
        super().__init__(convert_charrefs=True)
        self.doc = document
        self.figures_dir = figures_dir
        self._fig_counter = 0
        self._fmt_stack: List[str] = []
        self._current_para = None
        self._pending_rows: List[list] = []
        self._current_row: Optional[list] = None
        self._cell_text: List[str] = []
        self._cell_is_header = False
        self._in_table = False
        self._table_caption = ""
        self._in_caption = False
        self._in_figcaption = False
        self._figcaption_text: List[str] = []

    # -- helpers ---------------------------------------------------------
    def _add_run(self, text: str):
        # o HTML de report.py é indentado em f-strings multi-linha só para
        # legibilidade do código-fonte — normaliza como um navegador faria
        # (todo run de espaço/quebra de linha vira um único espaço) antes
        # de decidir se cria parágrafo/run algum.
        text = re.sub(r"\s+", " ", text)
        if not text:
            return
        if self._in_caption:
            # <caption> é filho de <table>, então precisa ser checado ANTES
            # de _in_table (senão o texto da legenda cairia em _cell_text).
            self._table_caption += text
            return
        if self._in_table:
            self._cell_text.append(text)
            return
        if self._in_figcaption:
            self._figcaption_text.append(text)
            return
        if text == " ":
            # espaço puro entre tags de bloco (nenhum parágrafo aberto) não
            # deve criar um parágrafo vazio; espaço puro logo no início de
            # um parágrafo recém-aberto também é descartado (mesmo
            # comportamento de colapso de espaço em branco de um navegador).
            if self._current_para is None or not self._current_para.runs:
                return
        if self._current_para is None:
            self._current_para = self.doc.add_paragraph()
        run = self._current_para.add_run(text)
        if "b" in self._fmt_stack or "strong" in self._fmt_stack:
            run.bold = True
        if "i" in self._fmt_stack or "em" in self._fmt_stack:
            run.italic = True
        if "sup" in self._fmt_stack:
            run.font.superscript = True
        if "sub" in self._fmt_stack:
            run.font.subscript = True
        if "code" in self._fmt_stack:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

    def _flush_table(self):
        if self._table_caption.strip():
            p = self.doc.add_paragraph()
            r = p.add_run(self._table_caption.strip())
            r.bold = True
            r.font.size = Pt(10)
        n_cols = max((len(row) for row in self._pending_rows), default=0)
        if n_cols and self._pending_rows:
            table = self.doc.add_table(rows=len(self._pending_rows), cols=n_cols)
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(self._pending_rows):
                for ci, (is_header, text) in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(text)
                    run.font.size = Pt(9.5)
                    if is_header:
                        run.bold = True
        self.doc.add_paragraph()
        self._table_caption = ""
        self._pending_rows = []

    def _add_figure(self):
        self._fig_counter += 1
        if not self.figures_dir:
            return
        fpath = os.path.join(self.figures_dir, f"Figure_{self._fig_counter}.png")
        if not os.path.exists(fpath):
            return
        with Image.open(fpath) as im:
            w_px, _ = im.size
        # largura no documento: min(20cm, largura real da imagem a 300dpi) —
        # nunca estica ALÉM da resolução real disponível.
        width_cm = min(MIN_FIGURE_WIDTH_CM, w_px / 300 * 2.54)
        self.doc.add_picture(fpath, width=Cm(width_cm))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _flush_figcaption(self):
        # chamado em </figcaption> — a legenda só termina de ser coletada
        # DEPOIS de <img/> já ter sido processado (a imagem vem antes da
        # legenda na ordem do HTML), então a figura não pode inserir a
        # legenda sozinha em _add_figure().
        text = "".join(self._figcaption_text).strip()
        self._figcaption_text = []
        if not text:
            return
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(text)
        r.italic = True
        r.font.size = Pt(9)

    # -- HTMLParser overrides ---------------------------------------------
    def handle_starttag(self, tag, attrs):
        if tag == "h1":
            self._current_para = self.doc.add_heading(level=0)
        elif tag == "h2":
            self._current_para = self.doc.add_heading(level=1)
        elif tag == "h3":
            self._current_para = self.doc.add_heading(level=2)
        elif tag == "p":
            self._current_para = self.doc.add_paragraph()
        elif tag == "li":
            self._current_para = self.doc.add_paragraph(style="List Bullet")
        elif tag in self._INLINE_TAGS:
            self._fmt_stack.append(tag)
        elif tag == "table":
            self._in_table = True
            self._pending_rows = []
        elif tag == "caption":
            self._in_caption = True
        elif tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._cell_text = []
            self._cell_is_header = (tag == "th")
        elif tag == "figcaption":
            self._in_figcaption = True
        elif tag == "img":
            self._add_figure()
        elif tag == "br":
            if self._current_para is not None:
                self._current_para.add_run().add_break()

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "p", "li"):
            self._current_para = None
        elif tag in self._INLINE_TAGS:
            if tag in self._fmt_stack:
                self._fmt_stack.reverse()
                self._fmt_stack.remove(tag)
                self._fmt_stack.reverse()
        elif tag == "tr":
            if self._current_row is not None:
                self._pending_rows.append(self._current_row)
            self._current_row = None
        elif tag in ("td", "th"):
            text = "".join(self._cell_text).strip()
            if self._current_row is not None:
                self._current_row.append((self._cell_is_header, text))
        elif tag == "table":
            self._flush_table()
            self._in_table = False
        elif tag == "caption":
            self._in_caption = False
        elif tag == "figcaption":
            self._in_figcaption = False
            self._flush_figcaption()

    def handle_data(self, data):
        self._add_run(data)


_BODY_RE = re.compile(r"<body[^>]*>(.*)</body>", re.S)
_DIV_BLOCK_RE = re.compile(r'<div class="(?:note|admin-banner)"[^>]*>.*?</div>', re.S)


def extract_docx_ready_body(full_html: str) -> str:
    """A partir do HTML AUTOCONTIDO final que report.py grava em disco
    (`<!doctype html>...<body>...</body></html>`), isola só o conteúdo que
    deve virar o .docx: o `<body>`, com o banner do Administrador/nota de
    aviso (`<div class="note">`/`<div class="admin-banner">`) removido —
    esse aviso é sobre o PROCESSO de geração, não conteúdo do manuscrito.
    Evita ter que montar um "body_html" paralelo em cada gerador."""
    m = _BODY_RE.search(full_html)
    body = m.group(1) if m else full_html
    return _DIV_BLOCK_RE.sub("", body)


def html_body_to_docx(body_html: str, out_path: str, doc_title: Optional[str] = None,
                       figures_dir: Optional[str] = None) -> str:
    """Converte um fragmento de corpo HTML (SEM <html>/<head>/banner admin)
    já montado por report.py em um .docx e grava em `out_path`. `doc_title`
    vira a propriedade de metadado "title" do arquivo docx (não um <h1>
    extra — o <h1> do próprio body_html já vira o título visível)."""
    document = Document()
    if doc_title:
        from html import unescape as _unescape
        plain_title = _unescape(re.sub(r"<[^>]+>", "", doc_title)).strip()
        # metadado core "title" do OOXML tem limite de 255 caracteres
        # (títulos deste projeto às vezes passam disso) — trunca só o
        # METADADO; o <h1> visível no corpo do documento fica completo.
        document.core_properties.title = plain_title[:255]
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    converter = _DocxHTMLConverter(document, figures_dir=figures_dir)
    converter.feed(body_html)
    converter.close()

    if os.path.dirname(out_path):
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
    document.save(out_path)
    return out_path


# --------------------------------------------------------------------------
# Tabela isolada em DOCX próprio (Table_N.docx)
# --------------------------------------------------------------------------

def table_to_docx(caption: str, headers: Sequence[str], rows: Sequence[Sequence[str]],
                   out_path: str) -> str:
    """Grava UMA tabela como seu próprio arquivo .docx (mesmos dados usados
    inline no relatório principal), pronta para anexo/revisão isolada."""
    document = Document()
    document.core_properties.title = re.sub(r"<[^>]+>", "", caption)[:255]
    p = document.add_paragraph()
    r = p.add_run(caption)
    r.bold = True

    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for ci, h in enumerate(headers):
        run = table.cell(0, ci).paragraphs[0].add_run(h)
        run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.cell(ri, ci).paragraphs[0].add_run(str(val))

    if os.path.dirname(out_path):
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
    document.save(out_path)
    return out_path


# --------------------------------------------------------------------------
# Carta de apresentação (cover letter)
# --------------------------------------------------------------------------

_AUTHOR_BLOCK = (
    "Cesar Rogerio Leal do Amaral (corresponding author), Cesar Rogerio do Amaral\n"
    "Núcleo de Genética Molecular Ambiental e Astrobiologia, Departamento de Biofísica "
    "e Biometria, Instituto de Biologia Roberto Alcantara Gomes, Universidade do Estado "
    "do Rio de Janeiro (UERJ)\n"
    "Avenida São Francisco Xavier, 524, Maracanã, Pavilhão Haroldo Lisboa da Cunha, "
    "CEP 20550-900, Rio de Janeiro, RJ, Brazil\n"
    "cesar.amaral@uerj.br"
)


def _new_letter_document() -> Document:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    p = document.add_paragraph(datetime.now().strftime("%Y-%m-%d"))
    p = document.add_paragraph()
    for line in _AUTHOR_BLOCK.split("\n"):
        p.add_run(line)
        p.add_run().add_break()
    document.add_paragraph()
    return document


def build_cover_letter_astrobiology(title: str, article_type: str, out_path: str) -> str:
    """Carta de apresentação para a Astrobiology (SAGE). Sem checklist
    formal publicado pela revista (ver comentário do módulo), mas inclui a
    exigência REAL e específica da SAGE de declarar assistência de
    terceiros (incluindo IA) na redação/edição, também na carta — não só em
    Acknowledgments."""
    document = _new_letter_document()
    document.add_paragraph("Dear Editor,")
    document.add_paragraph(
        f'Please find enclosed our manuscript, "{title}", which we would like to submit to '
        f"Astrobiology as a {article_type}."
    )
    document.add_paragraph(
        "[ADMIN, completar antes de enviar: 2-3 frases sobre a contribuição do estudo e por "
        "que ele se encaixa no escopo da Astrobiology — ver Aims and Scope da revista.]"
    )
    document.add_paragraph(
        "In accordance with SAGE's authorship policy on writing assistance and third-party "
        "submissions, we disclose that this manuscript's software tooling and successive "
        "drafts were developed with the assistance of an AI coding/writing assistant (Claude, "
        "Anthropic), under the full direction, review, and final approval of the human "
        "author(s); all scientific claims, data, and conclusions are the authors' own "
        "responsibility. [ADMIN: confirmar que esta frase reflete com precisão o processo real "
        "antes de enviar — ver também a seção de divulgação de uso de IA no paper de software "
        "da JOSS deste mesmo projeto.]"
    )
    document.add_paragraph(
        "The authors declare no competing interests. This manuscript is not under "
        "consideration elsewhere. [ADMIN, confirmar/editar.]"
    )
    document.add_paragraph("Thank you for your consideration.")
    document.add_paragraph()
    document.add_paragraph("Sincerely,")
    document.add_paragraph("Cesar Rogerio Leal do Amaral (corresponding author)")

    if os.path.dirname(out_path):
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
    document.save(out_path)
    return out_path


def build_cover_letter_plosone(title: str, contribution_summary: str, out_path: str) -> str:
    """Carta de apresentação para a PLOS ONE, seguindo os 6 itens REAIS
    exigidos por journals.plos.org/plosone/s/submission-guidelines
    (verificado 2026-08-09): contribuição do estudo, relação com a
    literatura, tipo de artigo, interações prévias com a PLOS, editores
    acadêmicos sugeridos, revisores a excluir. Limite de 1 página; NUNCA
    pedir isenção/redução de taxa aqui (vai no sistema de submissão)."""
    document = _new_letter_document()
    document.add_paragraph("Dear Editor,")
    document.add_paragraph(
        f'We wish to submit our manuscript, "{title}", as a Research Article for consideration '
        f"by PLOS ONE."
    )
    document.add_paragraph("Study contribution:").runs[0].bold = True
    document.add_paragraph(contribution_summary)
    document.add_paragraph("Relation to previously published work:").runs[0].bold = True
    document.add_paragraph(
        "[ADMIN, completar: 2-3 frases situando o estudo na literatura existente sobre a "
        "concentração prebiótica em fumarolas hidrotermais — ver Introduction do manuscrito.]"
    )
    document.add_paragraph("Article type: Research Article.")
    document.add_paragraph(
        "Prior interactions with PLOS regarding this manuscript: none. [ADMIN, editar se "
        "houver submissão anterior/desk-reject a mencionar.]"
    )
    document.add_paragraph(
        "Suggested Academic Editors: [ADMIN, completar — sugerir 1-3 nomes com expertise em "
        "astrobiologia/origem da vida/hidrotermalismo, se conhecidos.]"
    )
    document.add_paragraph(
        "Opposed reviewers: none. [ADMIN, editar se houver.]"
    )
    document.add_paragraph(
        "In accordance with PLOS's authorship policies, we disclose that this manuscript's "
        "software tooling and successive drafts were developed with the assistance of an AI "
        "coding/writing assistant (Claude, Anthropic), under the full direction, review, and "
        "final approval of the human author(s). [ADMIN: confirmar antes de enviar.]"
    )
    document.add_paragraph(
        "We confirm that this manuscript has not been published previously and is not under "
        "consideration in another journal. The authors have no competing interests to declare. "
        "[ADMIN, confirmar/editar.]"
    )
    document.add_paragraph("Thank you for your consideration.")
    document.add_paragraph()
    document.add_paragraph("Sincerely,")
    document.add_paragraph("Cesar Rogerio Leal do Amaral (corresponding author)")

    if os.path.dirname(out_path):
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
    document.save(out_path)
    return out_path


# --------------------------------------------------------------------------
# Arquivo suplementar (.zip) — conveniência do autor; ver ressalva SAGE no
# comentário do módulo sobre preferir arquivos individuais na submissão real
# --------------------------------------------------------------------------

def build_supplementary_zip(out_zip_path: str, items: Sequence[tuple], readme_text: str) -> str:
    """`items`: sequência de (caminho_no_disco, nome_dentro_do_zip). Grava
    também um README.txt (a partir de `readme_text`) na raiz do zip."""
    if os.path.dirname(out_zip_path):
        os.makedirs(os.path.dirname(out_zip_path), exist_ok=True)
    with zipfile.ZipFile(out_zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("README.txt", readme_text)
        for src_path, arcname in items:
            if src_path and os.path.exists(src_path):
                zf.write(src_path, arcname)
    return out_zip_path
